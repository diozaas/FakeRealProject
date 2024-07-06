import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
from docx import Document
import openpyxl
from openpyxl.chart import BarChart, Reference
import matplotlib.pyplot as plt
import pickle

# Load data
hoax_data = pd.read_excel('HoaxData.xlsx')
real_data = pd.read_excel('SatkerData.xlsx')

# Add labels
hoax_data['label'] = 0  # 0 for fake news
real_data['label'] = 1  # 1 for real news

# Combine datasets
data = pd.concat([hoax_data, real_data], ignore_index=True)

# Preprocess data
X_title = data['Title']
X_description = data['Description']
X_both = data['Title'] + " " + data['Description']
y = data['label']

# Function to evaluate models
def evaluate_models(X, y):
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
    
    # Vectorize data
    vectorizer = TfidfVectorizer()
    X_train_vec = vectorizer.fit_transform(X_train)
    X_test_vec = vectorizer.transform(X_test)
    
    # Define models
    models = {
        'Naive Bayes': MultinomialNB(),
        'Logistic Regression': LogisticRegression(max_iter=200),
        'Random Forest': RandomForestClassifier()
    }
    
    results = {}
    for model_name, model in models.items():
        model.fit(X_train_vec, y_train)
        y_pred = model.predict(X_test_vec)
        accuracy = accuracy_score(y_test, y_pred)
        report = classification_report(y_test, y_pred)
        cm = confusion_matrix(y_test, y_pred)
        results[model_name] = {'accuracy': accuracy, 'report': report, 'confusion_matrix': cm}

    random_forest_model = RandomForestClassifier()
    random_forest_model.fit(X_train_vec, y_train)
    # Save the model and vectorizer
    with open('random_forest_model.pkl', 'wb') as model_file:
        pickle.dump(random_forest_model, model_file)

    with open('tfidf_vectorizer.pkl', 'wb') as vectorizer_file:
        pickle.dump(vectorizer, vectorizer_file)
    
    return results

# Evaluate models with different feature sets
results_title = evaluate_models(X_title, y)
results_description = evaluate_models(X_description, y)
results_both = evaluate_models(X_both, y)

# Create a Word document to save results
doc = Document()
doc.add_heading('Comparison of Fake News Detection Models', 0)

def add_results_to_doc(doc, feature_set_name, results):
    doc.add_heading(feature_set_name, level=1)
    for model_name, metrics in results.items():
        doc.add_heading(model_name, level=2)
        doc.add_paragraph(f'Accuracy: {metrics["accuracy"]:.2f}')
        doc.add_heading('Classification Report:', level=3)
        doc.add_paragraph(metrics['report'])
        doc.add_heading('Confusion Matrix:', level=3)
        cm_table = doc.add_table(rows=3, cols=3)
        cm_table.style = 'Table Grid'
        cm_table.cell(0, 0).text = ''
        cm_table.cell(0, 1).text = 'Predicted Fake'
        cm_table.cell(0, 2).text = 'Predicted Real'
        cm_table.cell(1, 0).text = 'Actual Fake'
        cm_table.cell(2, 0).text = 'Actual Real'
        cm_table.cell(1, 1).text = str(metrics['confusion_matrix'][0, 0])
        cm_table.cell(1, 2).text = str(metrics['confusion_matrix'][0, 1])
        cm_table.cell(2, 1).text = str(metrics['confusion_matrix'][1, 0])
        cm_table.cell(2, 2).text = str(metrics['confusion_matrix'][1, 1])

add_results_to_doc(doc, 'Title Only', results_title)
add_results_to_doc(doc, 'Description Only', results_description)
add_results_to_doc(doc, 'Title and Description', results_both)

# Save the document
doc.save('FakeRealResult.docx')

# Prepare data for chart
accuracy_data = {
    'Feature Set': ['Title Only', 'Description Only', 'Title and Description'],
    'Naive Bayes': [results_title['Naive Bayes']['accuracy'], results_description['Naive Bayes']['accuracy'], results_both['Naive Bayes']['accuracy']],
    'Logistic Regression': [results_title['Logistic Regression']['accuracy'], results_description['Logistic Regression']['accuracy'], results_both['Logistic Regression']['accuracy']],
    'Random Forest': [results_title['Random Forest']['accuracy'], results_description['Random Forest']['accuracy'], results_both['Random Forest']['accuracy']]
}

accuracy_df = pd.DataFrame(accuracy_data)

# Create an Excel file with chart
with pd.ExcelWriter('FakeRealResultChart.xlsx') as writer:
    accuracy_df.to_excel(writer, sheet_name='Accuracy', index=False)
    workbook = writer.book
    worksheet = writer.sheets['Accuracy']
    
    # Create a bar chart
    chart = BarChart()
    chart.type = "col"
    chart.style = 10
    chart.title = "Model Accuracy Comparison"
    chart.y_axis.title = 'Accuracy'
    chart.x_axis.title = 'Feature Set'
    
    data = Reference(worksheet, min_col=2, min_row=1, max_col=4, max_row=4)
    categories = Reference(worksheet, min_col=1, min_row=2, max_row=4)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(categories)
    
    worksheet.add_chart(chart, "F2")

# Plot accuracy comparison chart using matplotlib
plt.figure(figsize=(10, 6))
accuracy_df.plot(x='Feature Set', kind='bar', rot=0)
plt.title('Model Accuracy Comparison')
plt.ylabel('Accuracy')
plt.xlabel('Feature Set')
plt.legend(loc='lower right')
plt.tight_layout()
plt.savefig('FakeRealResultChart.png')

